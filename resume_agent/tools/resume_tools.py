"""Resume domain tools - wrap pure domain functions as BaseTool instances.

Each tool handles file I/O (reading content from disk) and delegates
the actual resume logic to ``resume_agent.domain``.
"""

from __future__ import annotations

import hashlib
import json
from pathlib import Path
from typing import Tuple

from resume_agent.core.tools.base import BaseTool, ToolResult
from resume_agent.domain.job_matcher import format_match_report, match_job
from resume_agent.domain.resume_linter import format_lint_report, lint_resume
from resume_agent.domain.resume_parser import extract_sections, json_resume_to_text
from resume_agent.domain.resume_validator import (
    format_validation_report,
    validate_resume,
)
from resume_agent.domain.resume_writer import (
    markdown_to_html,
    markdown_to_json_resume,
    markdown_to_plain_text,
)

# ---------------------------------------------------------------------------
# ResumeParserTool
# ---------------------------------------------------------------------------


class ResumeParserTool(BaseTool):
    """Parse resume files into structured data."""

    name = "resume_parse"
    description = (
        "Parse one resume file into normalized text for analysis/editing workflows. "
        "Supports .pdf, .docx, .md, .txt, and JSON Resume input. "
        "Returns full extracted content plus detected section previews."
    )
    parameters = {
        "path": {
            "type": "string",
            "description": "Absolute or workspace-relative path to a resume file (.pdf/.docx/.md/.txt/.json).",
            "required": True,
        },
    }

    def __init__(self, workspace_dir: str = "."):
        self.workspace_dir = Path(workspace_dir).resolve()

    async def execute(self, path: str) -> ToolResult:
        try:
            file_path = self._resolve_path(path)
            if not file_path.exists():
                return ToolResult(success=False, output="", error=f"File not found: {path}")

            suffix = file_path.suffix.lower()

            if suffix == ".pdf":
                content, metadata = await self._parse_pdf(file_path)
            elif suffix == ".docx":
                content, metadata = await self._parse_docx(file_path)
            elif suffix in [".md", ".txt"]:
                content = file_path.read_text(encoding="utf-8")
                metadata = {"lines": content.count("\n") + 1, "characters": len(content)}
            elif suffix == ".json":
                raw = file_path.read_text(encoding="utf-8")
                data = json.loads(raw)
                content, metadata = json_resume_to_text(data)
            else:
                return ToolResult(
                    success=False,
                    output="",
                    error=f"Unsupported file format: {suffix}. Supported: .pdf, .docx, .md, .txt, .json",
                )

            sections = extract_sections(content)

            output = f"=== Resume Content ===\n{content}\n\n=== Detected Sections ===\n"
            for section, text in sections.items():
                output += f"\n[{section}]\n{text[:200]}{'...' if len(text) > 200 else ''}\n"

            result = ToolResult(
                success=True,
                output=output,
                data={
                    "path": str(file_path),
                    "format": suffix,
                    "sections": list(sections.keys()),
                    "metadata": metadata,
                },
            )
            return result
        except Exception as e:
            return ToolResult(success=False, output="", error=str(e))

    async def _parse_pdf(self, path: Path) -> Tuple[str, dict]:
        try:
            import fitz
        except ImportError:
            raise ImportError("PyMuPDF not installed. Run: pip install pymupdf")
        doc = fitz.open(path)
        text_parts = [page.get_text() for page in doc]
        metadata = {"pages": len(doc), "title": doc.metadata.get("title", ""), "author": doc.metadata.get("author", "")}
        doc.close()
        return "\n".join(text_parts), metadata

    async def _parse_docx(self, path: Path) -> Tuple[str, dict]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        doc = Document(path)
        text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        metadata = {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}
        return "\n".join(text_parts), metadata

    def _resolve_path(self, path: str) -> Path:
        p = Path(path)
        return p if p.is_absolute() else self.workspace_dir / p


# ---------------------------------------------------------------------------
# ResumeWriterTool
# ---------------------------------------------------------------------------


class ResumeWriterTool(BaseTool):
    """Write/generate resume files in various formats."""

    name = "resume_write"
    requires_approval = True
    mutation_signature_fields = ("path", "content", "template")
    description = (
        "Write resume content to disk, with output format determined by file extension. "
        "Supported outputs: .md, .txt, .json, .html. "
        "Does not generate .pdf/.docx directly; export those via a separate conversion step."
    )
    parameters = {
        "path": {
            "type": "string",
            "description": "Target output path; extension must be one of .md/.txt/.json/.html.",
            "required": True,
        },
        "content": {
            "type": "string",
            "description": "Resume content to write. Markdown is recommended as the source format.",
            "required": True,
        },
        "template": {
            "type": "string",
            "description": "HTML theme name (applies only when path ends with .html). Default: modern.",
            "default": "modern",
        },
    }

    def __init__(self, workspace_dir: str = "."):
        self.workspace_dir = Path(workspace_dir).resolve()
        self._preview_manager = None

    async def execute(self, path: str, content: str, template: str = "modern") -> ToolResult:
        try:
            file_path = self._resolve_path(path)
            file_path.parent.mkdir(parents=True, exist_ok=True)
            suffix = file_path.suffix.lower()

            if suffix == ".md":
                output_content = content
            elif suffix == ".txt":
                output_content = markdown_to_plain_text(content)
            elif suffix == ".json":
                output_content = markdown_to_json_resume(content)
            elif suffix == ".html":
                css = self._load_template_css(template)
                output_content = markdown_to_html(content, css=css)
            else:
                return ToolResult(
                    success=False,
                    output="",
                    error=f"Unsupported output format: {suffix}. Supported: .md, .txt, .json, .html",
                )

            fingerprint = hashlib.sha1(output_content.encode("utf-8", "ignore")).hexdigest()[:12]
            changed = True
            if file_path.exists() and file_path.is_file():
                try:
                    changed = file_path.read_text(encoding="utf-8") != output_content
                except Exception:
                    changed = True

            if self._preview_manager is not None:
                self._preview_manager.add(path, output_content, file_path)
                return ToolResult(
                    success=True,
                    output=f"Successfully wrote resume to {path} ({len(output_content)} characters)",
                    data={
                        "preview": True,
                        "pending_path": path,
                        "path": str(file_path),
                        "format": suffix,
                        "size": len(output_content),
                        "changed": changed,
                        "fingerprint_after": fingerprint,
                    },
                )

            file_path.write_text(output_content, encoding="utf-8")
            return ToolResult(
                success=True,
                output=f"Successfully wrote resume to {path} ({len(output_content)} characters)",
                data={
                    "path": str(file_path),
                    "format": suffix,
                    "size": len(output_content),
                    "changed": changed,
                    "fingerprint_after": fingerprint,
                },
            )
        except Exception as e:
            return ToolResult(success=False, output="", error=str(e))

    def _load_template_css(self, template: str) -> str | None:
        """Try to load CSS from core templates; return None to use domain fallback."""
        try:
            from resume_agent.core.templates import load_template_css

            return load_template_css(template)
        except Exception:
            return None

    def _resolve_path(self, path: str) -> Path:
        p = Path(path)
        return p if p.is_absolute() else self.workspace_dir / p


# ---------------------------------------------------------------------------
# ResumeLinterTool
# ---------------------------------------------------------------------------


class ResumeLinterTool(BaseTool):
    """Lint a resume for structure, formatting, and keyword quality."""

    name = "lint_resume"
    description = (
        "Score one resume (0-100) and report actionable quality issues. "
        "Breakdown includes formatting, completeness, keywords, and structure. "
        "Optionally accepts a job description to evaluate keyword alignment."
    )
    parameters = {
        "path": {
            "type": "string",
            "description": "Absolute or workspace-relative path to a text-based resume file.",
            "required": True,
        },
        "job_description": {
            "type": "string",
            "description": "Optional JD text used to score relevance and keyword overlap.",
        },
        "lang": {
            "type": "string",
            "description": "Language routing override: auto, en, or zh. Default: auto.",
            "default": "auto",
        },
    }

    def __init__(self, workspace_dir: str = "."):
        self.workspace_dir = Path(workspace_dir).resolve()

    async def execute(self, path: str, job_description: str = "", lang: str = "auto") -> ToolResult:
        try:
            file_path = self._resolve_path(path)
            if not file_path.exists():
                return ToolResult(success=False, output="", error=f"File not found: {path}")

            content = file_path.read_text(encoding="utf-8")
            if not content.strip():
                return ToolResult(success=False, output="", error=f"File is empty: {path}")

            result = lint_resume(content, job_description, lang=lang)
            output = format_lint_report(result)

            return ToolResult(
                success=True,
                output=output,
                data={
                    "overall_score": result.overall_score,
                    "sections": {
                        "formatting": {"score": result.formatting[0], "issues": result.formatting[1]},
                        "completeness": {"score": result.completeness[0], "issues": result.completeness[1]},
                        "keywords": {"score": result.keywords[0], "issues": result.keywords[1]},
                        "structure": {"score": result.structure[0], "issues": result.structure[1]},
                    },
                    "suggestions": result.suggestions,
                },
            )
        except Exception as e:
            return ToolResult(success=False, output="", error=str(e))

    def _resolve_path(self, path: str) -> Path:
        p = Path(path)
        return p if p.is_absolute() else self.workspace_dir / p


# ---------------------------------------------------------------------------
# JobMatcherTool
# ---------------------------------------------------------------------------


class JobMatcherTool(BaseTool):
    """Match a resume against a job description and identify gaps."""

    name = "job_match"
    description = (
        "Compare one resume against one job description to quantify fit and identify gaps. "
        "Returns match score, matched/missing keywords, extracted requirements, and edit suggestions. "
        "Primary input is job_text; this tool does not fetch job_url content by itself."
    )
    parameters = {
        "resume_path": {
            "type": "string",
            "description": "Absolute or workspace-relative path to the resume file to analyze.",
            "required": True,
        },
        "job_text": {
            "type": "string",
            "description": "Full job description text to compare against (recommended input).",
        },
        "job_url": {
            "type": "string",
            "description": "Compatibility field only. If provided without job_text, call fails; use web_read first.",
        },
    }

    def __init__(self, workspace_dir: str = "."):
        self.workspace_dir = Path(workspace_dir).resolve()

    async def execute(self, resume_path: str, job_text: str = "", job_url: str = "") -> ToolResult:
        try:
            file_path = self._resolve_path(resume_path)
            if not file_path.exists():
                return ToolResult(success=False, output="", error=f"Resume not found: {resume_path}")

            resume_content = file_path.read_text(encoding="utf-8")
            if not resume_content.strip():
                return ToolResult(success=False, output="", error=f"Resume is empty: {resume_path}")

            if not job_text.strip() and not job_url.strip():
                return ToolResult(success=False, output="", error="Provide either job_text or job_url")

            jd = job_text.strip()
            if not jd and job_url.strip():
                return ToolResult(
                    success=False,
                    output="",
                    error="job_url provided but fetching not supported in this tool. "
                    "Use web_read tool first, then pass the text via job_text.",
                )

            result = match_job(resume_content, jd)
            output = format_match_report(result)

            return ToolResult(
                success=True,
                output=output,
                data={
                    "match_score": result.match_score,
                    "location_score": result.location_score,
                    "skill_score": result.skill_score,
                    "yoe_score": result.yoe_score,
                    "company_experience_score": result.company_experience_score,
                    "keyword_score": result.keyword_score,
                    "requirement_score": result.requirement_score,
                    "semantic_score": result.semantic_score,
                    "score_breakdown": result.score_breakdown,
                    "skill_breakdown": result.skill_breakdown,
                    "semantic_evidence": result.semantic_evidence,
                    "backend_info": result.backend_info,
                    "quick_insights": result.quick_insights,
                    "next_step": result.next_step,
                    "matched_keywords": sorted(result.matched_keywords),
                    "missing_keywords": sorted(result.missing_keywords),
                    "suggestions": result.suggestions,
                    "requirements": result.requirements,
                },
            )
        except Exception as e:
            return ToolResult(success=False, output="", error=str(e))

    def _resolve_path(self, path: str) -> Path:
        p = Path(path)
        return p if p.is_absolute() else self.workspace_dir / p


# ---------------------------------------------------------------------------
# ResumeValidatorTool
# ---------------------------------------------------------------------------


class ResumeValidatorTool(BaseTool):
    """Validate resume content for completeness and format correctness."""

    name = "resume_validate"
    description = (
        "Validate a resume file for structural completeness and common quality/format issues. "
        "Returns pass/fail plus detailed errors and warnings for follow-up fixes."
    )
    parameters = {
        "path": {
            "type": "string",
            "description": "Absolute or workspace-relative path to the resume file to validate.",
            "required": True,
        },
    }

    def __init__(self, workspace_dir: str = "."):
        self.workspace_dir = Path(workspace_dir).resolve()

    async def execute(self, path: str) -> ToolResult:
        try:
            file_path = self._resolve_path(path)
            if not file_path.exists():
                return ToolResult(success=False, output="", error=f"File not found: {path}")

            content = file_path.read_text(encoding="utf-8")
            suffix = file_path.suffix.lower()

            result = validate_resume(content, file_format=suffix)
            output = format_validation_report(path, result)

            return ToolResult(
                success=True,
                output=output,
                data={
                    "valid": result.valid,
                    "errors": result.errors,
                    "warnings": result.warnings,
                    "format": suffix,
                },
            )
        except Exception as e:
            return ToolResult(success=False, output="", error=str(e))

    def _resolve_path(self, path: str) -> Path:
        p = Path(path)
        return p if p.is_absolute() else self.workspace_dir / p

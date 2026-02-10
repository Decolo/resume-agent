"""Resume parser tool - extract structured data from various resume formats."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Dict, Tuple
from .base import BaseTool, ToolResult


class ResumeParserTool(BaseTool):
    """Parse resume files into structured data."""

    name = "resume_parse"
    description = """Parse a resume file (PDF, DOCX, MD, TXT, JSON) and extract its content.
Returns structured text that can be analyzed and modified.
Supported formats: .pdf, .docx, .md, .txt, .json"""
    parameters = {
        "path": {
            "type": "string",
            "description": "Path to the resume file",
            "required": True,
        },
    }

    def __init__(self, workspace_dir: str = "."):
        self.workspace_dir = Path(workspace_dir).resolve()
        # Cache: path -> (mtime, parsed_result)
        self._cache: Dict[str, Tuple[float, ToolResult]] = {}

    async def execute(self, path: str) -> ToolResult:
        try:
            file_path = self._resolve_path(path)
            if not file_path.exists():
                return ToolResult(success=False, output="", error=f"File not found: {path}")

            # Check cache based on file modification time
            current_mtime = file_path.stat().st_mtime
            cache_key = str(file_path)

            if cache_key in self._cache:
                cached_mtime, cached_result = self._cache[cache_key]
                if cached_mtime == current_mtime:
                    # Cache hit - file hasn't changed
                    return cached_result

            suffix = file_path.suffix.lower()

            if suffix == ".pdf":
                content, metadata = await self._parse_pdf(file_path)
            elif suffix == ".docx":
                content, metadata = await self._parse_docx(file_path)
            elif suffix in [".md", ".txt"]:
                content, metadata = await self._parse_text(file_path)
            elif suffix == ".json":
                content, metadata = await self._parse_json(file_path)
            else:
                return ToolResult(
                    success=False,
                    output="",
                    error=f"Unsupported file format: {suffix}. Supported: .pdf, .docx, .md, .txt, .json",
                )

            # Try to extract sections
            sections = self._extract_sections(content)

            output = f"=== Resume Content ===\n{content}\n\n=== Detected Sections ===\n"
            for section, text in sections.items():
                output += f"\n[{section}]\n{text[:200]}{'...' if len(text) > 200 else ''}\n"

            result = ToolResult(
                success=True,
                output=output,
                data={
                    "path": str(file_path),
                    "format": suffix,
                    "sections": list(sections.keys()),
                    "metadata": metadata,
                },
            )

            # Store in cache
            self._cache[cache_key] = (current_mtime, result)

            return result

        except Exception as e:
            return ToolResult(success=False, output="", error=str(e))

    async def _parse_pdf(self, path: Path) -> Tuple[str, dict]:
        """Parse PDF file using PyMuPDF."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF not installed. Run: pip install pymupdf")

        doc = fitz.open(path)
        text_parts = []
        metadata = {
            "pages": len(doc),
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
        }

        for page in doc:
            text_parts.append(page.get_text())

        doc.close()
        return "\n".join(text_parts), metadata

    async def _parse_docx(self, path: Path) -> Tuple[str, dict]:
        """Parse DOCX file using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = Document(path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }

        return "\n".join(text_parts), metadata

    async def _parse_text(self, path: Path) -> Tuple[str, dict]:
        """Parse plain text or Markdown file."""
        content = path.read_text(encoding="utf-8")
        metadata = {
            "lines": content.count("\n") + 1,
            "characters": len(content),
        }
        return content, metadata

    async def _parse_json(self, path: Path) -> Tuple[str, dict]:
        """Parse JSON resume (JSON Resume format)."""
        content = path.read_text(encoding="utf-8")
        data = json.loads(content)
        
        # Convert JSON to readable text
        text_parts = []
        
        if "basics" in data:
            basics = data["basics"]
            text_parts.append(f"# {basics.get('name', 'Unknown')}")
            text_parts.append(f"{basics.get('label', '')}")
            text_parts.append(f"Email: {basics.get('email', '')}")
            text_parts.append(f"Phone: {basics.get('phone', '')}")
            if "location" in basics:
                loc = basics["location"]
                text_parts.append(f"Location: {loc.get('city', '')}, {loc.get('region', '')}")
            text_parts.append(f"\n{basics.get('summary', '')}")

        if "work" in data:
            text_parts.append("\n## Work Experience")
            for job in data["work"]:
                text_parts.append(f"\n### {job.get('position', '')} at {job.get('company', '')}")
                text_parts.append(f"{job.get('startDate', '')} - {job.get('endDate', 'Present')}")
                text_parts.append(job.get("summary", ""))
                for highlight in job.get("highlights", []):
                    text_parts.append(f"- {highlight}")

        if "education" in data:
            text_parts.append("\n## Education")
            for edu in data["education"]:
                text_parts.append(f"\n### {edu.get('studyType', '')} in {edu.get('area', '')}")
                text_parts.append(f"{edu.get('institution', '')}")
                text_parts.append(f"{edu.get('startDate', '')} - {edu.get('endDate', '')}")

        if "skills" in data:
            text_parts.append("\n## Skills")
            for skill in data["skills"]:
                keywords = ", ".join(skill.get("keywords", []))
                text_parts.append(f"- {skill.get('name', '')}: {keywords}")

        metadata = {"format": "json_resume", "sections": list(data.keys())}
        return "\n".join(text_parts), metadata

    def _extract_sections(self, content: str) -> Dict[str, str]:
        """Extract common resume sections from content."""
        sections = {}
        
        # Common section headers
        section_patterns = [
            (r"(?i)(summary|objective|profile|about)", "summary"),
            (r"(?i)(experience|work\s*history|employment)", "experience"),
            (r"(?i)(education|academic|qualification)", "education"),
            (r"(?i)(skills|technical\s*skills|competencies)", "skills"),
            (r"(?i)(projects|portfolio)", "projects"),
            (r"(?i)(certifications?|licenses?)", "certifications"),
            (r"(?i)(awards?|honors?|achievements?)", "awards"),
            (r"(?i)(publications?|papers?)", "publications"),
            (r"(?i)(languages?)", "languages"),
            (r"(?i)(references?)", "references"),
        ]

        lines = content.split("\n")
        current_section = "header"
        current_content = []

        for line in lines:
            matched = False
            for pattern, section_name in section_patterns:
                if re.match(pattern, line.strip()):
                    if current_content:
                        sections[current_section] = "\n".join(current_content).strip()
                    current_section = section_name
                    current_content = []
                    matched = True
                    break
            
            if not matched:
                current_content.append(line)

        if current_content:
            sections[current_section] = "\n".join(current_content).strip()

        return sections

    def _resolve_path(self, path: str) -> Path:
        p = Path(path)
        if p.is_absolute():
            return p
        return self.workspace_dir / p
